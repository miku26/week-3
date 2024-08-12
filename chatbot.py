import streamlit as st
from streamlit import _bottom
from groq import Groq
import matplotlib.pyplot as plt
import seaborn as sns
import speech_recognition as sr
import json
import numpy as np
import regex as re
import base64
import sqlite3
from datetime import datetime
import os
from pandasai import SmartDataframe
from pandasai.exceptions import NoCodeFoundError
import shutil
import time
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import AzureAIDocumentIntelligenceLoader
import groq
import PyPDF2
import docx
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_groq import ChatGroq
from langchain.prompts import ChatPromptTemplate
from langchain.docstore.document import Document
from dotenv import load_dotenv
import pandas as pd
import io
from pandasai.llm.base import LLM
import time
from groq import RateLimitError

# Load environment variables
load_dotenv()
user_defined_path = os.getcwd()
# Define constants
FAISS_PATH = "faiss_index"  # Path to store FAISS index
DB_NAME = 'chat_sessions.db'  # SQLite database name

# Initialize environment and clients
GROQ_API_KEY = os.getenv('GROQ_API_KEY')
client = Groq(api_key=GROQ_API_KEY, base_url='https://api.groq.com')
embedding_function = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
model = ChatGroq(temperature=0, model_name="mixtral-8x7b-32768")

# Define the prompt template for RAG
PROMPT_TEMPLATE = """
Answer the question based only on the following context:

{context}

---

Answer the question based on the above context: {question}
"""

# File reading functions
def read_pdf(content):
    """Read and extract text from a PDF file."""
    pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
    return " ".join(page.extract_text() for page in pdf_reader.pages)

def read_text(content):
    """Read and decode text from a text file."""
    return content.decode('utf-8')

def read_docx(content):
    """Read and extract text from a DOCX file."""
    doc = docx.Document(io.BytesIO(content))
    return " ".join(paragraph.text for paragraph in doc.paragraphs)


def read_csv(content):
    return pd.read_csv(content)

def read_xlsx(content):
    return pd.read_excel(io.BytesIO(content))


def create_smart_dataframe(data):
    llm = ChatGroq(model_name="mixtral-8x7b-32768", api_key = os.environ["GROQ_API_KEY"])    
    return SmartDataframe(data, config={
        "llm": llm,
        "enable_cache": True,
        "conversation_memory": True,
        "verbose": True,
        "save_charts": True,
        "use_error_correction_framework":True,
        "_max_retries":6,
        "allow_dangerous_code":True
    })


def display_data_preview(data):
    """Display a preview of the data."""
    st.write("Data preview:")
    st.dataframe(data.head())

def pandasai_query(query, sdf):
    if query:
        try:
            # Get response from SmartDataframe
            response = sdf.chat(str(query))
            
            # Check the type of response and handle accordingly
            if isinstance(response, str) and response.startswith("exports/charts/"):
                # Handle file path response
                with open(response, "rb") as img_file:
                    return img_file.read()
            elif isinstance(response, str):
                # Handle string response
                return response
            elif isinstance(response, pd.DataFrame):
                # Handle tabular response
                return response
            elif isinstance(response, plt.Figure):
                # Handle graphical response
                buf = io.BytesIO()
                response.savefig(buf, format="png")
                return buf.getvalue()
            else:
                # Handle unknown response type
                return f"Unsupported response type: {type(response)}"
        except NoCodeFoundError:
            return "No code found for the given query."
        except Exception as e:
            return f"An error occurred: {e}"
    return None

def read_file(file):
    """Read content from various file types."""
    content = file.read()
    file_readers = {
        "application/pdf": read_pdf,
        "text/": read_text,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": read_docx,
        "text/csv": read_csv,
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": read_xlsx

    }
    for file_type, reader_func in file_readers.items():
        if file.type.startswith(file_type):
            return reader_func(content) 
    raise ValueError(f"Unsupported file type: {file.type}")

# Database functions
def initialize_db(embedding_function, uploaded_file):
    """Initialize or reinitialize the FAISS database with uploaded file content."""
    if os.path.exists(FAISS_PATH):
        shutil.rmtree(FAISS_PATH)
        print("Removed existing FAISS index.")

    print("Creating a new database...")
    documents = []

    try:
        content = read_file(uploaded_file)
        print(f"Content of uploaded file (first 100 characters):")
        print(content[:100])
        print("---")
        documents.append(Document(page_content=content, metadata={"source": uploaded_file.name}))

    except Exception as e:
        print(f"Error reading uploaded file: {str(e)}")
    

    if not documents:
        print("No documents were successfully loaded. Cannot create database.")
        return None
    
    db = FAISS.from_documents(documents, embedding_function)
    db.save_local(FAISS_PATH)
    print("Documents added and database saved.")
    
    return db

def setup_sqlite_db():
    """Set up SQLite database for storing chat sessions and messages."""
    conn = sqlite3.connect(DB_NAME, check_same_thread=False)
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS sessions
                 (id INTEGER PRIMARY KEY AUTOINCREMENT, 
                  created_at TIMESTAMP)''')
    c.execute('''CREATE TABLE IF NOT EXISTS messages
                 (id INTEGER PRIMARY KEY AUTOINCREMENT, 
                  session_id INTEGER,
                  role TEXT,
                  content TEXT,
                  timestamp TIMESTAMP,
                  FOREIGN KEY (session_id) REFERENCES sessions(id))''')
    conn.commit()
    return conn, c

def create_new_session(cursor, conn):
    """Create a new chat session in the database."""
    cursor.execute("INSERT INTO sessions (created_at) VALUES (?)", (datetime.now(),))
    conn.commit()
    return cursor.lastrowid

def get_all_sessions_with_first_message(cursor):
    """Retrieve all chat sessions with their first user message."""
    cursor.execute("""
        SELECT s.id, m.content, s.created_at
        FROM sessions s
        LEFT JOIN messages m ON s.id = m.session_id AND m.role = 'user'
        WHERE m.id = (
            SELECT MIN(id)
            FROM messages
            WHERE session_id = s.id AND role = 'user'
        )
        ORDER BY s.created_at DESC
    """)
    return cursor.fetchall()

def get_session_messages(cursor, session_id):
    """Retrieve all messages for a given session."""
    cursor.execute("SELECT role, content FROM messages WHERE session_id = ? ORDER BY timestamp", (session_id,))
    messages = []
    for role, content in cursor.fetchall():
        try:
            # Try to parse as JSON first
            parsed_content = json.loads(content)
        except json.JSONDecodeError:
            # If it's not JSON, it might be a base64 encoded image
            if content.startswith('data:image') or (len(content) % 4 == 0 and re.match(r'^[A-Za-z0-9+/]+={0,2}$', content)):
                parsed_content = base64_to_image(content)
            else:
                parsed_content = content
        messages.append((role, parsed_content))
    return messages

def image_to_base64(image_bytes):
    return base64.b64encode(image_bytes).decode('utf-8')

def base64_to_image(base64_string):
    if base64_string.startswith('data:image'):
        # Remove the prefix if it exists
        base64_string = base64_string.split(',')[1]
    image_bytes = base64.b64decode(base64_string)
    return image_bytes

def add_message_to_db(cursor, conn, session_id, role, content):
    """Add a new message to the database."""
    if isinstance(content, bytes):  # Check if content is image bytes
        content = image_to_base64(content)
    else:
        content = json.dumps(content)
    cursor.execute("INSERT INTO messages (session_id, role, content, timestamp) VALUES (?, ?, ?, ?)",
              (session_id, role, content, datetime.now()))
    conn.commit()

def get_most_recent_user_message(cursor, session_id):
    """Get the most recent user message for a given session."""
    cursor.execute("""
        SELECT content
        FROM messages
        WHERE session_id = ? AND role = 'user'
        ORDER BY timestamp DESC
        LIMIT 1
    """, (session_id,))
    result = cursor.fetchone()
    return result[0] if result else "New Chat"

# AI and NLP functions
def voice_to_text():
    """Convert voice input to text using speech recognition."""
    r = sr.Recognizer()
    try:
        with sr.Microphone() as source:
            st.info("Microphone initialized. Listening... Speak now!")
            print("Microphone initialized. Listening...")  # Console debug print
            
            # Adjust for ambient noise
            r.adjust_for_ambient_noise(source, duration=1)
            
            try:
                audio = r.listen(source, timeout=5, phrase_time_limit=5)
                st.info("Audio captured. Processing speech...")
                print("Audio captured. Processing...")  # Console debug print
            except sr.WaitTimeoutError:
                st.warning("No speech detected. Listening stopped.")
                print("No speech detected.")  # Console debug print
                return None
            
        text = r.recognize_google(audio)
        print(f"Recognized text: {text}")  # Console debug print
        return text
    except sr.RequestError as e:
        st.error(f"Could not request results from Google Speech Recognition service; {e}")
        print(f"RequestError: {e}")  # Console debug print
    except sr.UnknownValueError:
        st.error("Google Speech Recognition could not understand audio")
        print("UnknownValueError")  # Console debug print
    except Exception as e:
        st.error(f"An unexpected error occurred: {str(e)}")
        print(f"Unexpected error: {str(e)}")  # Console debug print
    return None

def is_general_query(input_text):
    """Check if the input is a general query or greeting."""
    general_keywords = ['hello', 'hi', 'hey', 'how are you', 'what\'s up']
    return any(keyword in input_text.lower() for keyword in general_keywords)

def get_groq_response(input_text):
    """Get a response from the Groq API for general queries."""
    try:
        chat_completion = client.chat.completions.create(
            messages=[
                {
                    "role": "system",
                    "content": "You are a helpful assistant."
                },
                {
                    "role": "user",
                    "content": input_text
                }
            ],
            model="mixtral-8x7b-32768",
            temperature=0.5,
            max_tokens=1024,
            top_p=1,
            stop=None,
            stream=False
        )
        return chat_completion.choices[0].message.content
    except groq.APIError as e:
        if e.code == 502:
            return "I'm sorry, but the service is currently unavailable. Please try again later."
        else:
            return f"An API error occurred: {str(e)}"
    except Exception as e:
        return f"An unexpected error occurred: {str(e)}"
    
# UI functions
def setup_page():
    """Set up the Streamlit page configuration and custom CSS."""
    st.set_page_config(page_title="AI Chatbot", page_icon="🤖", layout="wide")
    st.markdown("""
    <style>
        .stTextInput > div > div > input {
            border-top-right-radius: 0;
            border-bottom-right-radius: 0;
        }
        .stButton > button {
            border-top-left-radius: 0;
            border-bottom-left-radius: 0;
            height: 42px;
        }
    </style>
    """, unsafe_allow_html=True)

def create_main_layout():
    """Create the main layout for the chat interface."""
    main_col1, main_col2, main_col3 = st.columns([1, 6, 1])
    return main_col2

def create_sidebar(conn, cursor):
    """Create the sidebar with chat history and controls."""
    with st.sidebar:
        st.image('image.png', width=200, use_column_width=True)
        
        if st.button("New Chat"):
            st.session_state.current_session_id = create_new_session(cursor, conn)
            st.rerun()
        
        current_message = get_most_recent_user_message(cursor, st.session_state.current_session_id)
        st.write(f"Prompt: {current_message[:50]}..." if len(current_message) > 50 else current_message)
        
        sessions = get_all_sessions_with_first_message(cursor)
        for session_id, first_message, created_at in sessions:
            if session_id != st.session_state.current_session_id:
                button_label = f"{first_message[:30]}..." if first_message else "Empty Chat"
                if st.button(button_label, key=f"session_{session_id}"):
                    st.session_state.current_session_id = session_id
                    st.rerun()
        
        if st.button("Clear All History"):
            cursor.execute("DELETE FROM messages")
            cursor.execute("DELETE FROM sessions")
            conn.commit()
            st.session_state.current_session_id = create_new_session(cursor, conn)
            st.rerun()

def RAG_response(messages, db, model):
    """Generate a response using Retrieval-Augmented Generation (RAG)."""
    print("RAG_response function called")
    print("Messages received:", messages)
    
    user_messages = [msg for msg in messages if msg['role'] == 'user']
    
    if not user_messages:
        return "No user messages found in the conversation."
    
    query_text = user_messages[-1]['content']
    print("Query text:", query_text)
    
    if not query_text.strip():
        return "The last user message is empty. Please provide a question or statement."
    
    results = db.similarity_search_with_score(query_text, k=1)  # Reduced to 1 document
    
    if len(results) == 0:
        return "I couldn't find any relevant information to answer your question."
    
    context_text = results[0][0].page_content
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,  # Further reduced
        chunk_overlap=100,
        length_function=len,
    )
    context_chunks = text_splitter.split_text(context_text)
    
    # Limit the number of chunks to reduce token usage
    max_chunks = 3
    context_chunks = context_chunks[:max_chunks]
    
    combined_context = "\n\n---\n\n".join(context_chunks)
    
    prompt_template = ChatPromptTemplate.from_template(PROMPT_TEMPLATE)
    prompt = prompt_template.format(context=combined_context, question=query_text)
    
    max_retries = 3
    retry_delay = 10  # seconds
    
    for attempt in range(max_retries):
        try:
            response_text = model.invoke(prompt)
            combined_response = response_text.content
            
            # Summarize only if the response is very long
            if len(combined_response.split()) > 500:
                summarization_prompt = ChatPromptTemplate.from_template(
                    "Summarize the following text in about 500 words: {text}"
                )
                summary_prompt = summarization_prompt.format(text=combined_response)
                summary_response = model.invoke(summary_prompt)
                final_response = summary_response.content
            else:
                final_response = combined_response

            sources = [results[0][0].metadata.get("source", None)]
            formatted_response = f"Response: {final_response}\nSource: {sources[0]}"
            
            return formatted_response
        
        except RateLimitError as e:
            if attempt < max_retries - 1:
                print(f"Rate limit reached. Retrying in {retry_delay} seconds...")
                time.sleep(retry_delay)
                retry_delay *= 2  # Exponential backoff
            else:
                return "I'm sorry, but I'm currently experiencing high demand. Please try again in a few minutes."
        
        except Exception as e:
            return f"An unexpected error occurred: {str(e)}"

    return "Failed to generate a response after multiple attempts. Please try again later."



def handle_file_upload(uploaded_file, db):
    if uploaded_file and not st.session_state.file_processed:
        if uploaded_file.type == "text/csv":
            data = read_csv(uploaded_file)
            st.session_state.smart_df = create_smart_dataframe(data)
            st.session_state.uploaded_file_type = "csv"
            print("smart df :", st.session_state.smart_df)
            st.success("CSV file is loaded. You can now query the data.")
            display_data_preview(data)
        else:
            db = initialize_db(embedding_function, uploaded_file)
            if db:
                st.success("document is uploaded.")
                st.session_state.db = db
                st.session_state.uploaded_file_type = uploaded_file.type
            else:
                st.error("Failed to update the database with the uploaded file.")
        
        st.session_state.file_processed = True
        st.session_state.file_uploader_key += 1
    
    return db

def handle_user_input(user_input, db, conn, cursor, chat_container, response_type):
    if user_input:
        st.session_state.messages.append({"role": "user", "content": user_input})
        add_message_to_db(cursor, conn, st.session_state.current_session_id, "user", user_input)
        
        with chat_container.chat_message("user"):
            st.markdown(user_input)

        if response_type == "File_Query":
            if st.session_state.uploaded_file_type == "csv":
                # Handle CSV query
                response = pandasai_query(str(user_input), st.session_state.smart_df)
            elif st.session_state.uploaded_file_type in ["application/pdf", "text/plain", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
                # Handle PDF, TXT, DOCX query
                response = RAG_response(st.session_state.messages, db, model)
            else:
                print("file type is ",st.session_state.uploaded_file_type)
                response = "Unsupported file type for querying."
        else:
            # Normal query response
            response = get_groq_response(user_input)

        # Handle the response (this part remains largely unchanged)
        if isinstance(response, str):
            st.session_state.messages.append({"role": "assistant", "content": response})
            with chat_container.chat_message("assistant"):
                st.markdown(response)
            add_message_to_db(cursor, conn, st.session_state.current_session_id, "assistant", response)
        elif isinstance(response, pd.DataFrame):
            st.session_state.messages.append({"role": "assistant", "content": str(response)})
            with chat_container.chat_message("assistant"):
                st.dataframe(response)
            add_message_to_db(cursor, conn, st.session_state.current_session_id, "assistant", str(response))
        elif isinstance(response, bytes):
            st.session_state.messages.append({"role": "assistant", "content": "Here's the chart:"})
            with chat_container.chat_message("assistant"):
                st.image(response)
            add_message_to_db(cursor, conn, st.session_state.current_session_id, "assistant", response)
        else:
            st.session_state.messages.append({"role": "assistant", "content": f"Unsupported response type: {type(response)}"})
            with chat_container.chat_message("assistant"):
                st.write(f"Unsupported response type: {type(response)}")
            add_message_to_db(cursor, conn, st.session_state.current_session_id, "assistant", str(response))
    

def process_input(user_input, uploaded_file, db, conn, cursor, chat_container, response_type):
    """Process user input, including file uploads and text/voice messages."""
    db = handle_file_upload(uploaded_file, db)
    handle_user_input(user_input, db, conn, cursor, chat_container, response_type)
    return db

def main():
    """Main function to run the Streamlit app."""
    setup_page()
    conn, cursor = setup_sqlite_db()
    
    # Initialize session state variables
    if "current_session_id" not in st.session_state:
        st.session_state.current_session_id = create_new_session(cursor, conn)
    
    if "file_processed" not in st.session_state:
        st.session_state.file_processed = False
    
    if "file_uploader_key" not in st.session_state:
        st.session_state.file_uploader_key = 0
    
    if "db" not in st.session_state:
        st.session_state.db = None
    
    if "messages" not in st.session_state:
        st.session_state.messages = [{"role": "system", "content": "You are a helpful assistant."}]
    
    
    if "current_session_id" not in st.session_state:
        st.session_state.current_session_id = create_new_session(cursor, conn)
    
    if "file_processed" not in st.session_state:
        st.session_state.file_processed = False
    
    if "smart_df" not in st.session_state:  
        st.session_state.smart_df = None

    if "uploaded_file_type" not in st.session_state:
        st.session_state.uploaded_file_type = None
    
    if "file_uploader_key" not in st.session_state:
        st.session_state.file_uploader_key = 0
    
    messages = get_session_messages(cursor, st.session_state.current_session_id)
    st.session_state.messages = [{"role": "system", "content": "You are a helpful assistant."}] + \
                                [{"role": role, "content": content} for role, content in messages]
    

    
    create_sidebar(conn, cursor)
    
    st.write("Welcome to your friendly AI chatbot! You can type or use voice input.")
    
    chat_container = st.container()
    
    with chat_container:
        for message in st.session_state.messages[1:]:
            with st.chat_message(message["role"]):
                content = message["content"]
                if isinstance(content, bytes):
                    st.image(content)
                elif isinstance(content, str):
                    if content.startswith("data:image"):
                        # It's a base64 encoded image
                        st.image(base64_to_image(content))
                    else:
                        st.markdown(content)
                elif isinstance(content, pd.DataFrame):
                    st.dataframe(content)
                else:
                    st.write(content)
    
    input_container = st.container()
    with input_container:
        if "voice_input" not in st.session_state:
            st.session_state.voice_input = ""
        response_type = st.radio("Choose response type:", ("File_Query", "Normal_Query"), horizontal=True)
        st.markdown('<div class="input-container">', unsafe_allow_html=True)
        col1, col2, col3 = _bottom.columns([6, 1, 1])
        with col1:
            user_input = st.chat_input(placeholder="Type your message here...", key="user_input")
        with col2:
            mic_button = st.button("🎤", key="mic_button")
        with col3:
            uploaded_file = st.file_uploader("Upload file", type=["txt", "pdf", "docx","csv"], key=f"file_uploader", label_visibility="collapsed")   
            


    css = '''
<style>
    /* Style the file uploader container */
    [data-testid='stFileUploader'] {
        width: 38px;
        height: 38px;
        overflow: hidden;
        position: relative;
    }

    /* Hide the text but keep the input functional */
    [data-testid='stFileUploader'] div {
        opacity: 0;
    }

    [data-testid='stFileUploader'] input[type="file"] {
        opacity: 0;
        position: absolute;
        width: 100%;
        height: 100%;
        top: 0;
        left: 0;
        cursor: pointer;
    }

    /* Create a new visual element for the icon */
    [data-testid='stFileUploader']::before {
        content: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' width='24' height='24' fill='%23000000'%3E%3Cpath d='M9 16h6v-6h4l-7-7-7 7h4v6zm-4 2h14v2H5v-2z'/%3E%3C/svg%3E");
        position: absolute;
        top: 50%;
        left: 50%;
        transform: translate(-50%, -50%);
        width: 24px;
        height: 24px;
        pointer-events: none;
    }

    /* Remove hover effects */
    [data-testid='stFileUploader']:hover::before {
        content: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' width='24' height='24' fill='%23000000'%3E%3Cpath d='M9 16h6v-6h4l-7-7-7 7h4v6zm-4 2h14v2H5v-2z'/%3E%3C/svg%3E");
    }

    /* Hide any tooltip */
    [data-testid='stFileUploader'] [title] {
        position: relative;
    }

    [data-testid='stFileUploader'] [title]:hover::after {
        content: none;
    }
</style>
'''

    st.markdown(css, unsafe_allow_html=True)

    if mic_button:
        st.info("Microphone button pressed. Initializing voice input...")
        print("Microphone button pressed.")  # Console debug print
        user_input = voice_to_text()
        if user_input is not None:
            st.info(f"Voice input received: {user_input}")
            print(f"Voice input received: {user_input}")  # Console debug print
            st.session_state.db = process_input(user_input, None, st.session_state.db, conn, cursor, chat_container)
        else:
            st.warning("No valid voice input received.")
            print("No valid voice input received.")  # Console debug print
        st.rerun()
    
    if uploaded_file or user_input:
        st.session_state.db = process_input(user_input, uploaded_file, st.session_state.db, conn, cursor, chat_container, response_type)
        if uploaded_file:
            st.session_state.file_processed = False

    conn.close()


if __name__ == "__main__":
    main()
